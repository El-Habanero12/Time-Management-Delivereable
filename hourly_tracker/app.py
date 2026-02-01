from __future__ import annotations

import os
import sys
import threading
import traceback
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Iterable, List, Optional

from PIL import Image, ImageDraw
from docx import Document
import pystray

import json

from hourly_tracker.paths import get_appdata_dir, get_default_expenses_path, get_docs_dir, is_test_profile
from hourly_tracker.first_run import ensure_user_files_exist


def _default_state_dir() -> Path:
    return get_appdata_dir()


def _default_data_dir() -> Path:
    return get_docs_dir()


def ensure_dir(p: Path) -> Path:
    path = Path(p)
    path.mkdir(parents=True, exist_ok=True)
    return path


@dataclass
class Config:
    interval_minutes: int = 60
    snooze_minutes: int = 10
    dismiss_snooze_minutes: int = 10
    pause_minutes: int = 60
    catch_up_max_hours: int = 8

    llm_enabled: bool = False
    llm_model: str = "llama3.1:8b"
    llm_timeout_seconds: int = 20

    no_network_mode: bool = True

    state_dir: Path = field(default_factory=_default_state_dir)
    data_dir: Path = field(default_factory=_default_data_dir)
    log_path: Path | None = None
    report_path: Path | None = None
    state_path: Path | None = None
    learned_rules_path: Path | None = None
    config_path: Path | None = None
    log_lock_path: Path | None = None
    reflections_dir: Path | None = None
    expenses_path: Path | None = field(default_factory=get_default_expenses_path)

    reflection_enabled: bool = True
    reflection_time_local: str = "23:30"

    _PATH_FIELDS = [
        "state_dir",
        "data_dir",
        "log_path",
        "report_path",
        "state_path",
        "learned_rules_path",
        "config_path",
        "log_lock_path",
        "reflections_dir",
        "expenses_path",
    ]

    def _coerce_path_fields(self) -> None:
        """Ensure all path-like attributes are Path instances."""
        for name in self._PATH_FIELDS:
            value = getattr(self, name, None)
            if value is None:
                continue
            if not isinstance(value, Path):
                try:
                    setattr(self, name, Path(value))
                except Exception:
                    pass

    def resolve_paths(self) -> "Config":
        self._coerce_path_fields()

        state_base = ensure_dir(Path(self.state_dir))
        data_base = ensure_dir(Path(self.data_dir))
        self.state_dir = state_base
        self.data_dir = data_base
        if self.log_path is None:
            self.log_path = data_base / "time_log.xlsx"
        if self.report_path is None:
            self.report_path = data_base / "report.html"
        if self.state_path is None:
            self.state_path = state_base / "state.json"
        if self.learned_rules_path is None:
            self.learned_rules_path = state_base / "learned_rules.json"
        if self.config_path is None:
            self.config_path = state_base / "config.json"
        if self.log_lock_path is None:
            self.log_lock_path = state_base / "time_log.lock"
        if self.reflections_dir is None:
            self.reflections_dir = data_base / "reflections"
        if self.expenses_path is not None and not isinstance(self.expenses_path, Path):
            self.expenses_path = Path(self.expenses_path)
        return self


def load_config() -> Config:
    cfg = Config().resolve_paths()
    if cfg.config_path and cfg.config_path.exists():
        try:
            data = json.loads(cfg.config_path.read_text(encoding="utf-8"))
            for key, val in data.items():
                if hasattr(cfg, key):
                    if key in Config._PATH_FIELDS and val is not None:
                        try:
                            setattr(cfg, key, Path(val))
                        except Exception:
                            setattr(cfg, key, val)
                    else:
                        setattr(cfg, key, val)
            cfg.resolve_paths()
        except Exception:
            pass
    # Always enforce hardcoded expenses path per user request.
    cfg.expenses_path = get_default_expenses_path()
    return cfg


def save_config(cfg: Config) -> None:
    cfg.resolve_paths()
    p = cfg.config_path
    if isinstance(p, str):
        p = Path(p)
        cfg.config_path = p
    if p:
        p.parent.mkdir(parents=True, exist_ok=True)
        payload = cfg.__dict__.copy()
        for k, v in list(payload.items()):
            if isinstance(v, Path):
                payload[k] = str(v)
        p.write_text(json.dumps(payload, indent=2), encoding="utf-8")

from hourly_tracker.dialogs import (
    CatchUpResult,
    PromptInput,
    PromptResult,
    TaskManagerResult,
    TkDialogRunner,
    ReflectionResult,
    SpendingResult,
    catch_up_dialog,
    error_dialog,
    prompt_dialog,
    reflection_dialog,
    spending_dialog,
    task_manager_dialog,
)
from hourly_tracker.excel_store import (
    DEFAULT_CATEGORIES,
    Entry,
    add_tasks,
    append_entry,
    append_reflection_index,
    append_to_expenses_workbook,
    ensure_workbook,
    log_task_event,
    upsert_daily_row,
    read_categories,
    read_tasks,
    update_task_fields,
)
from hourly_tracker.llm_ollama import detect_ollama, ollama_classify_category
from hourly_tracker.no_network import enforce_no_network
from hourly_tracker.notifications import Notifier
from hourly_tracker.scheduler import Scheduler
from hourly_tracker.state import StateStore
from hourly_tracker.tagging import CategorySuggester, Suggestion


@dataclass
class AppContext:
    cfg: Config
    dialog_runner: TkDialogRunner
    notifier: Notifier
    scheduler: Scheduler
    suggester: CategorySuggester
    tray_icon: Optional[pystray.Icon] = None


def _create_icon() -> Image.Image:
    size = 64
    img = Image.new("RGBA", (size, size), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    draw.ellipse((8, 8, size - 8, size - 8), fill=(52, 120, 246, 255))
    draw.rectangle((30, 16, 34, 36), fill=(255, 255, 255, 255))
    draw.rectangle((30, 36, 44, 40), fill=(255, 255, 255, 255))
    return img


def _log_event(cfg: Config, message: str) -> None:
    try:
        cfg.resolve_paths()
        log_path = cfg.state_dir / "app.log"
        log_path.parent.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().isoformat(timespec="seconds")
        with log_path.open("a", encoding="utf-8") as fh:
            fh.write(f"[{timestamp}] {message}\n")
    except Exception:
        pass


def _self_check_paths(cfg: Config) -> None:
    """Ensure key config paths are Path objects; coerce and log if not."""
    repaired: List[str] = []
    for name in Config._PATH_FIELDS:
        value = getattr(cfg, name, None)
        if value is None:
            continue
        if not isinstance(value, Path):
            try:
                setattr(cfg, name, Path(value))
                repaired.append(name)
            except Exception:
                _log_event(cfg, f"Could not coerce path field {name}: {value}")
    if repaired:
        _log_event(cfg, f"Coerced path fields to Path: {', '.join(repaired)}")
    cfg.resolve_paths()


def _available_categories(cfg: Config) -> List[str]:
    assert cfg.log_path is not None
    cats = read_categories(cfg.log_path, lock_path=cfg.log_lock_path)
    return cats or list(DEFAULT_CATEGORIES)


def _suggest_category(ctx: AppContext, activity: str) -> Optional[Suggestion]:
    return ctx.suggester.suggest(activity)


def _build_suggester(cfg: Config, categories: Iterable[str]) -> CategorySuggester:
    assert cfg.learned_rules_path is not None

    ollama_status = detect_ollama()
    llm_enabled = bool(cfg.llm_enabled and ollama_status.installed and cfg.llm_model in ollama_status.models)

    def _llm_classifier(activity: str, cats: List[str]) -> Optional[Suggestion]:
        if not llm_enabled:
            return None
        return ollama_classify_category(activity, cats, cfg.llm_model, cfg.llm_timeout_seconds)

    return CategorySuggester(
        learned_rules_path=cfg.learned_rules_path,
        categories=categories,
        llm_classifier=_llm_classifier,
        llm_enabled=llm_enabled,
    )


def _prompt_once(ctx: AppContext, prompt_type: str = "regular") -> None:
    categories = _available_categories(ctx.cfg)
    ctx.suggester.categories = categories
    open_tasks = read_tasks(ctx.cfg.log_path, status_filter="open", lock_path=ctx.cfg.log_lock_path) if ctx.cfg.log_path else []

    ctx.notifier.notify("Hourly Tracker", "Check-in due: what are you doing right now?")

    timestamp = datetime.now()

    def _dialog(root):
        # Start in full mode; quick entry is available via hotkeys/Enter.
        return prompt_dialog(
            root,
            categories,
            timestamp=timestamp,
            suggest_fn=lambda text: _suggest_category(ctx, text),
            tasks=open_tasks,
        )

    result = ctx.dialog_runner.run(_dialog)

    ctx.scheduler.mark_prompted(timestamp)

    if isinstance(result, PromptResult) and result.submitted and result.prompt_input:
        prompt_input = result.prompt_input
        # Suggestion based on final activity text.
        suggestion = _suggest_category(ctx, prompt_input.activity)
        chosen_category = prompt_input.category
        _ = suggestion
        # Learn from the user's explicit choice, especially overrides.
        ctx.suggester.learn_override(prompt_input.activity, chosen_category)
        _persist_prompt(ctx, prompt_input)
        _persist_task_updates(ctx, prompt_input)
        return

    # Dialog dismissed: auto-snooze.
    ctx.scheduler.snooze(ctx.cfg.dismiss_snooze_minutes)


def _persist_prompt(ctx: AppContext, prompt_input: PromptInput) -> None:
    entry = Entry(
        timestamp=prompt_input.timestamp,
        activity=prompt_input.activity,
        notes=prompt_input.notes,
        category=prompt_input.category,
        energy=prompt_input.energy,
        focus=prompt_input.focus,
        prompt_type=prompt_input.prompt_type,
    )
    assert ctx.cfg.log_path is not None
    categories = _available_categories(ctx.cfg)
    append_entry(ctx.cfg.log_path, entry, categories=categories, lock_path=ctx.cfg.log_lock_path)
    ctx.scheduler.mark_entry(prompt_input.timestamp)


def _persist_task_updates(ctx: AppContext, prompt_input: PromptInput) -> None:
    if not ctx.cfg.log_path:
        return

    new_tasks = prompt_input.new_tasks or []
    if new_tasks:
        add_tasks(ctx.cfg.log_path, new_tasks, lock_path=ctx.cfg.log_lock_path)

    worked_ids = set(prompt_input.worked_task_ids or [])
    completed_ids = set(prompt_input.completed_task_ids or [])

    minutes = int(prompt_input.task_minutes or 0)
    effort = int(prompt_input.task_effort or 3)
    could_be_faster = bool(prompt_input.task_could_be_faster)

    for task_id in worked_ids:
        log_task_event(
            ctx.cfg.log_path,
            task_id=task_id,
            action="worked",
            minutes=minutes,
            effort=effort,
            could_be_faster=could_be_faster,
            notes="",
            lock_path=ctx.cfg.log_lock_path,
        )

    for task_id in completed_ids:
        log_task_event(
            ctx.cfg.log_path,
            task_id=task_id,
            action="completed",
            minutes=minutes,
            effort=effort,
            could_be_faster=could_be_faster,
            notes="",
            lock_path=ctx.cfg.log_lock_path,
        )


def _catch_up(ctx: AppContext, hours_missed: int) -> None:
    categories = _available_categories(ctx.cfg)
    ctx.suggester.categories = categories

    ctx.notifier.notify("Hourly Tracker", f"You missed ~{hours_missed} hour(s). Let's catch up.")

    def _dialog(root):
        return catch_up_dialog(root, hours_missed=hours_missed, categories=categories)

    result = ctx.dialog_runner.run(_dialog)
    ctx.scheduler.mark_resume_handled()
    ctx.scheduler.mark_prompted()

    if not isinstance(result, CatchUpResult) or not result.submitted:
        ctx.scheduler.snooze(ctx.cfg.dismiss_snooze_minutes)
        return

    now = datetime.now()
    hours = max(1, int(result.hours))
    interval = timedelta(minutes=ctx.cfg.interval_minutes)

    entries: List[Entry] = []
    if result.split_entries and hours > 1:
        for i in range(hours):
            ts = now - interval * (hours - i)
            entries.append(
                Entry(
                    timestamp=ts,
                    activity=result.activity,
                    notes=result.notes,
                    category=result.category,
                    energy=result.energy,
                    focus=result.focus,
                    prompt_type="catch_up",
                    start_time=ts - interval,
                    end_time=ts,
                )
            )
    else:
        ts = now
        entries.append(
            Entry(
                timestamp=ts,
                activity=result.activity,
                notes=result.notes,
                category=result.category,
                energy=result.energy,
                focus=result.focus,
                prompt_type="catch_up",
                start_time=ts - interval * hours,
                end_time=ts,
            )
        )

    assert ctx.cfg.log_path is not None
    for entry in entries:
        append_entry(ctx.cfg.log_path, entry, categories=categories, lock_path=ctx.cfg.log_lock_path)
        ctx.scheduler.mark_entry(entry.timestamp)


def _log_spending(ctx: AppContext) -> None:
    def _dialog(root):
        return spending_dialog(root)

    result = ctx.dialog_runner.run(_dialog)
    if not isinstance(result, SpendingResult) or not result.submitted or not result.spending_input:
        return

    cfg = ctx.cfg
    cfg.resolve_paths()
    _, expenses_path = ensure_user_files_exist()
    cfg.expenses_path = expenses_path
    data = {
        "date": date.today(),
        "type": result.spending_input.entry_type,
        "payment_method": result.spending_input.payment_method,
        "amount": result.spending_input.amount,
        "notes": result.spending_input.notes,
    }
    try:
        upsert_daily_row(
            expenses_path,
            entry_date=data["date"],
            entry_type=data["type"],
            amount=data["amount"],
            method=data["payment_method"],
            notes=data["notes"],
        )
    except PermissionError:
        ctx.notifier.notify("Hourly Tracker", "Close Expenses.xlsx in Excel, then retry.")
        ctx.dialog_runner.run(lambda root: error_dialog(root, "Close Expenses.xlsx", "Please close Expenses.xlsx in Excel and try again."))
    except OSError as exc:
        ctx.notifier.notify("Hourly Tracker", "Expenses.xlsx is locked; close it and retry.")
        _log_event(cfg, f"log_spending locked error {exc}")
    except Exception:
        _log_event(cfg, f"log_spending error\n{traceback.format_exc()}")
        ctx.dialog_runner.run(lambda root: error_dialog(root, "Spending log failed", "Could not append to Expenses.xlsx. Check app.log for details."))


def _save_reflection(ctx: AppContext, reflection: ReflectionInput) -> Path:
    cfg = ctx.cfg
    cfg.resolve_paths()
    reflections_dir = Path(cfg.reflections_dir or (cfg.data_dir / "reflections"))
    reflections_dir.mkdir(parents=True, exist_ok=True)
    base_path = reflections_dir / f"{reflection.date_for.isoformat()}.docx"
    doc_path = base_path
    stamp = reflection.created_at.isoformat(timespec="seconds")

    try:
        doc = Document(doc_path) if doc_path.exists() else Document()
        loaded_existing = doc_path.exists()
    except Exception:
        # If the existing file is locked or corrupted, fall back to a timestamped filename.
        doc_path = reflections_dir / f"{reflection.date_for.isoformat()}_{reflection.created_at.strftime('%H%M%S')}.docx"
        _log_event(cfg, f"reflection doc fallback to {doc_path.name} due to load error")
        doc = Document()
        loaded_existing = False

    if not loaded_existing:
        doc.add_heading(f"Daily Reflection - {reflection.date_for.isoformat()}", level=0)
    else:
        doc.add_paragraph("")
        doc.add_paragraph(f"--- Added at {stamp} ---")

    # Preserve user-entered newlines by writing each line as its own paragraph
    content = reflection.text or " "
    for line in content.splitlines() or [" "]:
        doc.add_paragraph(line if line.strip() else " ")
    if reflection.tags:
        doc.add_paragraph(f"Tags: {reflection.tags}")
    doc.add_paragraph(f"Recorded at {stamp}")
    doc.save(doc_path)

    try:
        if cfg.log_path:
            append_reflection_index(
                cfg.log_path,
                cfg.log_lock_path,
                reflection.date_for.isoformat(),
                str(doc_path),
                stamp,
                None,
            )
    except Exception:
        _log_event(cfg, f"reflection index update failed\n{traceback.format_exc()}")

    return doc_path


def _handle_reflection(ctx: AppContext, date_for: date) -> None:
    def _dialog(root):
        return reflection_dialog(root, date_for)

    result = ctx.dialog_runner.run(_dialog)
    if isinstance(result, ReflectionResult) and result.submitted and result.reflection_input:
        try:
            _save_reflection(ctx, result.reflection_input)
        except Exception:
            _log_event(ctx.cfg, f"reflection save failed\n{traceback.format_exc()}")
            ctx.dialog_runner.run(
                lambda root: error_dialog(root, "Reflection save failed", "Could not save your reflection. It was logged in app.log.")
            )
    ctx.scheduler.mark_reflection_completed(date_for)


def _open_log(cfg: Config) -> None:
    assert cfg.log_path is not None
    try:
        os.startfile(cfg.log_path)  # type: ignore[attr-defined]
    except Exception:
        pass


def _open_task_manager(ctx: AppContext) -> None:
    tasks = read_tasks(ctx.cfg.log_path, status_filter="open", lock_path=ctx.cfg.log_lock_path) if ctx.cfg.log_path else []

    def _dialog(root):
        return task_manager_dialog(root, tasks)

    result = ctx.dialog_runner.run(_dialog)
    if not isinstance(result, TaskManagerResult) or not result.submitted:
        return

    if ctx.cfg.log_path:
        if result.added_tasks:
            add_tasks(ctx.cfg.log_path, result.added_tasks, lock_path=ctx.cfg.log_lock_path)
        for update in result.updated:
            update_task_fields(
                ctx.cfg.log_path,
                update.get("id"),
                title=update.get("title"),
                notes=update.get("notes"),
                lock_path=ctx.cfg.log_lock_path,
            )
        for task_id in result.completed_task_ids:
            log_task_event(
                ctx.cfg.log_path,
                task_id=task_id,
                action="completed",
                minutes=0,
                effort=3,
                could_be_faster=False,
                notes="",
                lock_path=ctx.cfg.log_lock_path,
            )


def _create_windows_shortcut(shortcut_path: Path, target_path: Path) -> None:
    try:
        import pythoncom  # type: ignore
        from win32com.client import Dispatch  # type: ignore
    except Exception as exc:  # pragma: no cover - platform specific
        raise RuntimeError("pywin32 is required to create shortcuts") from exc

    pythoncom.CoInitialize()
    try:
        shell = Dispatch("WScript.Shell")
        shortcut = shell.CreateShortcut(str(shortcut_path))
        shortcut.TargetPath = str(target_path)
        working_dir = target_path if target_path.is_dir() else target_path.parent
        shortcut.WorkingDirectory = str(working_dir)
        if target_path.suffix:
            shortcut.IconLocation = str(target_path)
        shortcut.Save()
    finally:  # pragma: no cover - platform specific
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def _create_shortcuts(ctx: AppContext, notify: bool = True, show_errors: bool = True) -> None:
    desktop = Path(os.environ.get("USERPROFILE") or Path.home()) / "Desktop"
    desktop.mkdir(parents=True, exist_ok=True)
    shortcuts = [
        (ctx.cfg.data_dir, "HourlyTracker Folder"),
        (ctx.cfg.log_path, "HourlyTracker Log"),
    ]
    created_any = False
    for target, name in shortcuts:
        if not target:
            continue
        shortcut_path = desktop / f"{name}.lnk"
        try:
            _create_windows_shortcut(shortcut_path, Path(target))
            created_any = True
        except Exception:
            _log_event(ctx.cfg, f"shortcut creation failed for {target}\n{traceback.format_exc()}")
            if show_errors:
                ctx.dialog_runner.run(
                    lambda root: error_dialog(root, "Shortcut error", f"Could not create shortcut for {target}. See app.log.")
                )
    if created_any:
        sentinel = ctx.cfg.state_dir / "shortcuts_created.flag"
        try:
            sentinel.write_text(datetime.now().isoformat(timespec="seconds"), encoding="utf-8")
        except Exception:
            _log_event(ctx.cfg, f"shortcut sentinel write failed\n{traceback.format_exc()}")
        if notify:
            ctx.notifier.notify("Hourly Tracker", "Desktop shortcuts created.")


def _ensure_shortcuts_once(ctx: AppContext) -> None:
    sentinel = ctx.cfg.state_dir / "shortcuts_created.flag"
    if sentinel.exists():
        return
    _create_shortcuts(ctx, notify=False, show_errors=False)


def _build_context() -> AppContext:
    cfg = load_config()

    # Ensure per-user working files are present in the profile docs dir.
    time_log_path, expenses_path = ensure_user_files_exist()
    cfg.data_dir = time_log_path.parent
    cfg.log_path = time_log_path
    cfg.expenses_path = expenses_path
    cfg.reflections_dir = cfg.data_dir / "reflections"

    _self_check_paths(cfg)
    save_config(cfg)

    if cfg.no_network_mode:
        enforce_no_network(allow_loopback=True)

    assert cfg.log_path is not None
    ensure_workbook(cfg.log_path, lock_path=cfg.log_lock_path)

    categories = _available_categories(cfg)
    suggester = _build_suggester(cfg, categories)

    dialog_runner = TkDialogRunner()
    notifier = Notifier()

    state_store = StateStore(cfg.state_path)
    scheduler = Scheduler(
        cfg=cfg,
        state_store=state_store,
        on_prompt=lambda: None,
        on_catch_up=lambda hours: None,
        on_reflection=None,
        tick_seconds=5.0,
    )

    # We need the context before scheduler callbacks run.
    ctx = AppContext(cfg=cfg, dialog_runner=dialog_runner, notifier=notifier, scheduler=scheduler, suggester=suggester)
    scheduler.on_prompt = lambda: _prompt_once(ctx)
    scheduler.on_catch_up = lambda hours: _catch_up(ctx, hours)
    scheduler.on_reflection = lambda day: _handle_reflection(ctx, day)

    threading.Thread(target=_ensure_shortcuts_once, args=(ctx,), daemon=True).start()
    return ctx


def run_tray_app() -> None:
    ctx = _build_context()

    tray_title = "Hourly Tracker (TEST)" if is_test_profile() else "Hourly Tracker"

    icon = pystray.Icon(
        "hourly-tracker",
        _create_icon(),
        tray_title,
        menu=pystray.Menu(
            pystray.MenuItem("Log now", lambda: _prompt_once(ctx, prompt_type="manual")),
            pystray.MenuItem("Snooze 10m", lambda: ctx.scheduler.snooze(ctx.cfg.snooze_minutes)),
            pystray.MenuItem("Pause 1h", lambda: ctx.scheduler.pause(ctx.cfg.pause_minutes)),
            pystray.MenuItem("Open log", lambda: _open_log(ctx.cfg)),
            pystray.MenuItem("Daily reflection", lambda: _handle_reflection(ctx, date.today())),
            pystray.MenuItem("Log today's spending", lambda: _log_spending(ctx)),
            pystray.MenuItem("Tasks...", lambda: _open_task_manager(ctx)),
            pystray.MenuItem("Create shortcuts", lambda: _create_shortcuts(ctx)),
            pystray.MenuItem("Quit", lambda: _quit(icon, ctx)),
        ),
    )

    ctx.tray_icon = icon
    ctx.scheduler.start()
    icon.run()


def _quit(icon: pystray.Icon, ctx: AppContext) -> None:
    try:
        ctx.scheduler.stop()
    finally:
        icon.stop()


def main() -> None:
    run_tray_app()


if __name__ == "__main__":
    # Support: python -m hourly_tracker.app
    main()
